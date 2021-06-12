import glob, os
import tika
tika.initVM()
from tika import parser
import re
import pdb
import pymongo
import string
import csv
import nltk
import shutil
import spacy
import nltk
from pdf2docx import Converter
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
stopset = set(stopwords.words('english'))
nlp = spacy.load('en_core_web_sm')
import heapq
from docx import Document
'''for file in glob.glob("*.pdf"):
    pdb.set_trace()
    docx_file = "./docxdownload_pdf_link/"+file[:-4]+".docx"
    cv = Converter(file)
    cv.convert(docx_file, start=0, end=None)
    cv.close()
    print('converted')'''
'''for file in glob.glob("*.pdf"):
    file_data = parser.from_file(file)
    text = file_data['content']
    try:
      data1 = text.strip('\n')
    except:
      print('not in english language')
    data  = data1.splitlines()
    data_2 = [x.strip() for x in data if x]
    data_3 = [x for x in data_2 if x != '']
    data_4 = [x for x in data_3 if len(x) != 1]
    data_5 = [x.replace('\xa0','') for x in data_4 if x]
    data_6 = [x.replace('\xad','') for x in data_5 if x]
    case_ref_data = []      
    data_pdf = []
    flag = False
    for i in data_6:
        string_1 =  "".join(i.lower().split())
        if flag:
           data_pdf.append(i)
        if 'judgment' in string_1:
             flag = True
        if 'order' in string_1:
            flag = True
        if 'corrigendum' in string_1:
            flag = True
    for i in data_pdf:
        string_1 =  "".join(i.lower().split())
        if 'v.' in string_1:
           case_ref_data.append(i)
        if 'vs.' in string_1:
           case_ref_data.append(i)
        if 'versus' in string_1:
           case_ref_data.append(i)
    remove_dup = (set(case_ref_data))
    final_out = list(remove_dup)
    print(file,final_out)'''
'''csv_file = open('cases_ref.csv','a')
    writer = csv.writer(csv_file)
    writer.writerow((file,final_out))'''
# coming paragrapgh which contains v.
flag = False
data_file = []
file_data = parser.from_file('0_2017_12_1501_19306_Judgement_07-Jan-2020.pdf')
text = file_data['content']
data_file.append(text)
data_pdf = []
flag = False
for i in data_file:
    string_1 =  "".join(i.lower().split())
    if flag:
       data_pdf.append(i)
    if 'judgment' in string_1:
        flag = True
    if 'order' in string_1:
        flag = True
    if 'corrigendum' in string_1:
        flag = True
# spllitting with '.' finding with v or v.
cases_ref = []
data_1 = [x.replace('\n\n','') for x in data_file]
data_2 = [x.splitlines() for x in data_1]
for i in data_2:
    for elem in i:
        if 'v.' in elem:
           if not elem.startswith(('v.','iv.')):
              cases_ref.append(elem)
              print(elem)
        if 'vs.' in elem:
            cases_ref.append(elem)
        if 'versus' in elem:
            cases_ref.append(elem)
for elm in cases_ref:
    print(elm)
'''for elem2 in cases_ref:
    elem1 = (elem2.split(','))
    for index,elem2 in enumerate(elem1):
        data_elem = []
        if 'v.' in elem2:
           if 'iv.' not in elem2:
              data_elem.append(elem2)
        if len(data_elem)>0:
           #data_elem1 = [i.split(' ') for i in data_elem]
           for elem4 in data_elem:
               print(elem4)
               document = Document()
               document.add_paragraph(elem4)
               document.save('docx_file.docx')
           document = Document('docx_file.docx')
           for para in document.paragraphs:
               data_1 = para.text.splitlines()
               for elems in data_1:
                   print(elems)'''
data1 = text.strip('\n')
data  = data1.splitlines()
data_2 = [x.strip() for x in data if x]
data_3 = [x for x in data_2 if x != '']
data_4 = [x for x in data_3 if len(x) != 1]
data_5 = [x.replace('\xa0','') for x in data_4 if x]
data_6 = [x.replace('\xad','') for x in data_5 if x]
data_pdf = []
flag = False
for i in data_6:
    string_1 =  "".join(i.lower().split())
    if flag:
       data_pdf.append(i)
    if 'judgment' in string_1:
        flag = True
    if 'order' in string_1:
        flag = True
    if 'corrigendum' in string_1:
        flag = True
case_ref_data = []
data_pdf_1 = [x for x in data_pdf if x != '']
#print(data_pdf_1)
'''for index,i in enumerate(data_pdf_1):
    string_1 =  "".join(i.lower().split())     
    if 'v.' in string_1:
       if not string_1.startswith(('v.','iv.')):
          case_ref_data.append(i)
    if 'vs.' in string_1:
       case_ref_data.append(i)
    if 'versus' in string_1:
        case_ref_data.append(i)'''
for elms in data_pdf_1:
    if elms[-1].isdigit():
        print(elms)
'''tokenized_models = [word_tokenize(i) for i in case_ref_data]
for i in tokenized_models:
    stop_models = [elem for elem in i if str(elem).lower() not in stopset]
    print(stop_models)
    #if 'Vs.' in i:
       #print(i)'''

'''doc = nlp('Court in Sree Ramakrishna Mining Company v. Commissioner')
for ent in doc.ents:
    print(ent.text,ent.label_)'''

'''import re
import textract
#read the content of pdf as text
text = textract.process('0_2017_12_1501_19306_Judgement_07-Jan-2020.pdf')
#use four space as paragraph delimiter to convert the text into list of paragraphs
text_1 = (text.decode("utf-8")) 
text_2 = (text_1.splitlines())
for elem in text_2:
    print(elem)'''

'''s = "abc1"
contains_digit = any(map(str.isdigit, s))
print(contains_digit)'''
